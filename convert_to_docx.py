"""
Markdown取扱説明書をWord (.docx) 形式に変換するスクリプト
画像も埋め込みます
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def add_formatted_text(paragraph, text):
    """テキスト内の **bold** と `code` をパースして追加"""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
        else:
            paragraph.add_run(part)

def md_to_docx(md_path, docx_path, title):
    doc = Document()
    
    # デフォルトフォント設定
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic UI'
    font.size = Pt(10.5)
    
    # ページ設定 A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # 空行
        if not line.strip():
            i += 1
            continue
        
        # テーブル行の収集
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line.strip())
            i += 1
            # テーブル終了チェック
            if i >= len(lines) or not (lines[i].strip().startswith('|') and lines[i].strip().endswith('|')):
                # テーブルを出力
                # セパレーター行を除去
                data_rows = [r for r in table_rows if not re.match(r'^\|[\s\-:]+\|$', r.replace('|', '|'))]
                filtered = []
                for r in data_rows:
                    cells = [c.strip() for c in r.split('|')[1:-1]]
                    if all(re.match(r'^[\-:]+$', c) for c in cells):
                        continue
                    filtered.append(cells)
                
                if filtered:
                    num_cols = len(filtered[0])
                    table = doc.add_table(rows=len(filtered), cols=num_cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    for ri, row_cells in enumerate(filtered):
                        for ci in range(min(num_cols, len(row_cells))):
                            cell = table.cell(ri, ci)
                            cell.text = ''
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(2)
                            p.paragraph_format.space_after = Pt(2)
                            add_formatted_text(p, row_cells[ci])
                            p.style.font.size = Pt(9.5)
                            if ri == 0:
                                for run in p.runs:
                                    run.bold = True
                    
                    doc.add_paragraph('')  # スペース
                
                in_table = False
                table_rows = []
            continue
        
        # 見出し
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        
        # 水平線
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            i += 1
            continue
        
        # 画像
        img_match = re.match(r'!\[.*?\]\((.*?)\)', line.strip())
        if img_match:
            img_rel = img_match.group(1)
            img_path = os.path.join(BASE_DIR, img_rel)
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(3.5))
                doc.add_paragraph('')  # スペース
            else:
                p = doc.add_paragraph(f'[画像: {img_rel}]')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        
        # 引用（> で始まる行）
        if line.startswith('> '):
            text = line[2:].strip()
            # 連続する引用行をまとめる
            while i + 1 < len(lines) and lines[i+1].strip().startswith('> '):
                i += 1
                text += '\n' + lines[i].strip()[2:].strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            # 背景色の代わりに左のマーカー
            add_formatted_text(p, '💬 ' + text)
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue
        
        # リスト項目
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', line)
        if list_match:
            indent = len(list_match.group(1))
            marker = list_match.group(2)
            content = list_match.group(3)
            
            if marker in ['-', '*']:
                p = doc.add_paragraph(style='List Bullet')
            else:
                p = doc.add_paragraph(style='List Number')
            
            if indent > 0:
                p.paragraph_format.left_indent = Cm(1.5 + indent * 0.3)
            
            add_formatted_text(p, content)
            i += 1
            continue
        
        # 通常の段落
        p = doc.add_paragraph()
        add_formatted_text(p, line.strip())
        i += 1
    
    doc.save(docx_path)
    print(f'[OK] Output: {docx_path}')


# --- オーナー・店長用 ---
md_to_docx(
    os.path.join(BASE_DIR, '取扱説明書_オーナー・店長用.md'),
    os.path.join(BASE_DIR, '取扱説明書_オーナー・店長用.docx'),
    'スマート出退勤 取扱説明書【オーナー・店長用】'
)

# --- スタッフ用 ---
md_to_docx(
    os.path.join(BASE_DIR, '取扱説明書_スタッフ用.md'),
    os.path.join(BASE_DIR, '取扱説明書_スタッフ用.docx'),
    'スマート出退勤 取扱説明書【スタッフ用】'
)

print('\n[DONE] All conversions completed!')
