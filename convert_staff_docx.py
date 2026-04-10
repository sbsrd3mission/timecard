"""
スタッフ用取扱説明書 -> Word (.docx) 変換スクリプト
- フォント: 游ゴシック (ゴシック体)
- 見出し+直後の内容が改ページで切れないよう keep_with_next を設定
- テーブルが途中で切れないよう keep_together を設定
- 画像の前後に適切な余白
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

GOTHIC_FONT = 'Yu Gothic'          # 游ゴシック
GOTHIC_FONT_HEADING = 'Yu Gothic'  # 見出しも同じゴシック


def set_font_gothic(run, size_pt=10.5, bold=False, color=None):
    """runのフォントを游ゴシックに統一する"""
    run.font.name = GOTHIC_FONT
    # 日本語フォントを明示的に指定
    run._element.rPr.rFonts.set(qn('w:eastAsia'), GOTHIC_FONT)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def set_keep_with_next(paragraph):
    """次の段落と同じページに収める（見出し用）"""
    paragraph.paragraph_format.keep_with_next = True


def set_keep_together(paragraph):
    """段落内で改ページしない"""
    paragraph.paragraph_format.keep_together = True


def add_formatted_text(paragraph, text, base_size=10.5, base_color=None):
    """テキスト内の **bold** と `code` をパースしてrunに追加（游ゴシック統一）"""
    parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            set_font_gothic(run, size_pt=base_size, bold=True, color=base_color)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
        else:
            run = paragraph.add_run(part)
            set_font_gothic(run, size_pt=base_size, color=base_color)


def apply_table_no_break(table):
    """テーブル全体がページをまたがないようにする"""
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), '1')
        trPr.append(cantSplit)


def style_heading(doc, level, text):
    """游ゴシックの見出しを追加し、keep_with_next を設定"""
    heading = doc.add_heading('', level=level)
    heading.clear()
    run = heading.add_run(text)
    
    if level == 0:
        set_font_gothic(run, size_pt=18, bold=True, color=RGBColor(0x4A, 0x3B, 0x32))
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        set_font_gothic(run, size_pt=14, bold=True, color=RGBColor(0x2E, 0x7D, 0x32))
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(6)
    elif level == 2:
        set_font_gothic(run, size_pt=12, bold=True, color=RGBColor(0x1A, 0x5C, 0x8A))
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(4)
    elif level == 3:
        set_font_gothic(run, size_pt=11, bold=True, color=RGBColor(0x55, 0x55, 0x55))
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(2)
    
    set_keep_with_next(heading)     # 見出しの直後で改ページしない
    set_keep_together(heading)
    return heading


def md_to_docx_staff(md_path, docx_path):
    doc = Document()

    # --- ドキュメント全体のデフォルトフォントをゴシックに ---
    style = doc.styles['Normal']
    style.font.name = GOTHIC_FONT
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), GOTHIC_FONT)

    # 箇条書きスタイルもフォント設定
    for style_name in ['List Bullet', 'List Number']:
        try:
            ls = doc.styles[style_name]
            ls.font.name = GOTHIC_FONT
            ls.font.size = Pt(10.5)
        except Exception:
            pass

    # A4 ページ設定
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_table = False
    table_rows = []
    # セクション状態管理（見出しの後の段落はすべて keep_with_next を連鎖させる）
    prev_was_heading = False

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # --- 空行 ---
        if not line.strip():
            prev_was_heading = False
            i += 1
            continue

        # --- テーブル ---
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line.strip())
            i += 1
            if i >= len(lines) or not (lines[i].strip().startswith('|') and lines[i].strip().endswith('|')):
                # セパレーター除去
                filtered = []
                for r in table_rows:
                    cells = [c.strip() for c in r.split('|')[1:-1]]
                    if all(re.match(r'^[\-:]+$', c) for c in cells):
                        continue
                    filtered.append(cells)

                if filtered:
                    num_cols = len(filtered[0])
                    table = doc.add_table(rows=len(filtered), cols=num_cols)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for ri, row_cells in enumerate(filtered):
                        for ci in range(min(num_cols, len(row_cells))):
                            cell = table.cell(ri, ci)
                            cell.text = ''
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(2)
                            p.paragraph_format.space_after = Pt(2)
                            set_keep_together(p)
                            cell_color = None
                            if ri == 0:
                                cell_color = RGBColor(0x33, 0x33, 0x33)
                            add_formatted_text(p, row_cells[ci],
                                               base_size=9.5,
                                               base_color=cell_color)
                            if ri == 0:
                                for run in p.runs:
                                    run.bold = True

                    # テーブル全体を改ページに強い設定に
                    apply_table_no_break(table)

                    sp = doc.add_paragraph('')
                    sp.paragraph_format.space_after = Pt(4)

                in_table = False
                table_rows = []
                prev_was_heading = False
            continue

        # --- 見出し ---
        if line.startswith('# '):
            style_heading(doc, 0, line[2:].strip())
            prev_was_heading = True
            i += 1
            continue

        if line.startswith('## '):
            style_heading(doc, 1, line[3:].strip())
            prev_was_heading = True
            i += 1
            continue

        if line.startswith('### '):
            style_heading(doc, 2, line[4:].strip())
            prev_was_heading = True
            i += 1
            continue

        if line.startswith('#### '):
            style_heading(doc, 3, line[5:].strip())
            prev_was_heading = True
            i += 1
            continue

        # --- 水平線 ---
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run('─' * 55)
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
            prev_was_heading = False
            i += 1
            continue

        # --- 画像 ---
        img_match = re.match(r'!\[.*?\]\((.*?)\)', line.strip())
        if img_match:
            img_rel = img_match.group(1)
            img_path = os.path.join(BASE_DIR, img_rel)
            # 画像はセクションの冒頭で使われることが多い -> keep_with_next
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            set_keep_with_next(p)
            if os.path.exists(img_path):
                run = p.add_run()
                run.add_picture(img_path, width=Inches(3.5))
            else:
                run = p.add_run(f'[画像: {img_rel}]')
                set_font_gothic(run, size_pt=9, color=RGBColor(0x88, 0x88, 0x88))
            prev_was_heading = False
            i += 1
            continue

        # --- 引用 (> ...) ---
        if line.startswith('> '):
            text = line[2:].strip()
            while i + 1 < len(lines) and lines[i + 1].strip().startswith('> '):
                i += 1
                text += '\n' + lines[i].strip()[2:].strip()

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.right_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            set_keep_together(p)
            if prev_was_heading:
                set_keep_with_next(p)  # 見出しの直後は次ともくっつける
            add_formatted_text(p, text, base_size=9.5,
                               base_color=RGBColor(0x44, 0x44, 0x88))
            prev_was_heading = False
            i += 1
            continue

        # --- リスト項目 ---
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', line)
        if list_match:
            indent = len(list_match.group(1))
            marker = list_match.group(2)
            content = list_match.group(3)

            if marker in ['-', '*']:
                p = doc.add_paragraph(style='List Bullet')
            else:
                p = doc.add_paragraph(style='List Number')

            if indent > 0:
                p.paragraph_format.left_indent = Cm(1.5 + indent * 0.3)

            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            set_keep_together(p)
            if prev_was_heading:
                set_keep_with_next(p)

            add_formatted_text(p, content)
            prev_was_heading = False
            i += 1
            continue

        # --- 通常の段落 ---
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        set_keep_together(p)
        if prev_was_heading:
            set_keep_with_next(p)
        add_formatted_text(p, line.strip())
        prev_was_heading = False
        i += 1

    doc.save(docx_path)
    print(f'[OK] Saved: {docx_path}')


# --- スタッフ用のみ変換 ---
md_to_docx_staff(
    os.path.join(BASE_DIR, 'staff_manual', '取扱説明書_スタッフ用.md')
    if os.path.exists(os.path.join(BASE_DIR, 'staff_manual', '取扱説明書_スタッフ用.md'))
    else os.path.join(BASE_DIR, '取扱説明書_スタッフ用.md'),
    os.path.join(BASE_DIR, '取扱説明書_スタッフ用.docx'),
)

print('\n[DONE] Staff manual converted successfully!')
